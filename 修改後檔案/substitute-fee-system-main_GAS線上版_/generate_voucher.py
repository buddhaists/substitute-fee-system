#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
代課費計算與管理系統 - 月度印領清冊 Word 產生器
功能：讀取網頁端導出的 JSON 備份檔，自動生成符合學校行政核銷格式的 Word 印領清冊 (.docx)
"""

import os
import json
import glob
from datetime import datetime

def install_and_import_docx():
    try:
        import docx
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
        return docx, Pt, Inches, WD_ALIGN_PARAGRAPH, OxmlElement, parse_xml, qn, nsdecls
    except ImportError:
        print("偵測到尚未安裝 python-docx 模組，正在嘗試為您自動安裝...")
        import subprocess
        import sys
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            import docx
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement, parse_xml
            from docx.oxml.ns import nsdecls, qn
            print("python-docx 安裝成功！\n")
            return docx, Pt, Inches, WD_ALIGN_PARAGRAPH, OxmlElement, parse_xml, qn, nsdecls
        except Exception as e:
            print(f"自動安裝失敗，請手動於終端機執行: pip install python-docx")
            print(f"錯誤訊息: {e}")
            sys.exit(1)

# 初始化 docx 相關元件
docx, Pt, Inches, WD_ALIGN_PARAGRAPH, OxmlElement, parse_xml, qn, nsdecls = install_and_import_docx()

def install_and_import_openpyxl():
    try:
        import openpyxl
        return openpyxl
    except ImportError:
        print("偵測到尚未安裝 openpyxl 模組，正在嘗試為您自動安裝...")
        import subprocess
        import sys
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
            import openpyxl
            print("openpyxl 安裝成功！\n")
            return openpyxl
        except Exception as e:
            print("自動安裝 openpyxl 失敗，請手動於終端機執行: pip install openpyxl")
            print(f"錯誤訊息: {e}")
            return None

# 初始化 openpyxl
openpyxl = install_and_import_openpyxl()

def set_cell_border(cell, **kwargs):
    """
    設定表格儲存格邊框
    kwargs 可以是 top, bottom, left, right
    值為 dictionary, 例如: top={'sz': 12, 'val': 'single', 'color': '000000'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_background(cell, fill_color):
    """設定儲存格背景顏色 (HEX)"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_roc_date(date_str):
    """將 YYYY-MM-DD 轉為 民國年/月/日"""
    if not date_str:
        return ""
    try:
        s = str(date_str)
        dt = datetime.strptime(s, "%Y-%m-%d")
        roc_year = dt.year - 1911  # 西元轉民國年
        return f"{roc_year}/{dt.month:02d}/{dt.day:02d}"
    except:
        return str(date_str)

def generate_excel(target_month, filtered_records):
    print("\n==================================================")
    print("           正在產生 Excel 試算表憑證...           ")
    print("==================================================")
    
    template_file = "印領清冊113.xlsx"
    if not os.path.exists(template_file):
        print(f"[錯誤] 找不到 Excel 範本檔案 '{template_file}'，無法產生 Excel 憑證！")
        return False
        
    if openpyxl is None:
        print("[錯誤] 未安裝 openpyxl，無法產生 Excel 憑證。")
        return False
        
    try:
        wb = openpyxl.load_workbook(template_file)
        
        # 解析月份 (YYYY-MM)
        year_str, month_str = target_month.split('-')
        year = int(year_str)
        month = int(month_str)
        roc_year = year - 1911
        
        # 1. 處理課堂鐘點費代課 (鐘點代課)
        # 排除整天導師費代課 (mentorDaily)
        period_records = [r for r in filtered_records if r.get('payMode', 'perPeriod') != 'mentorDaily']
        
        if period_records:
            # 尋找範本工作表 (預設用最新一期的格式，即 11506)
            source_sheet_name = '11506'
            if source_sheet_name not in wb.sheetnames:
                num_sheets = sorted([name for name in wb.sheetnames if name.isdigit()])
                source_sheet_name = num_sheets[-1] if num_sheets else wb.sheetnames[0]
                
            print(f"使用工作表 '{source_sheet_name}' 作為鐘點代課範本...")
            source_sheet = wb[source_sheet_name]
            
            # 建立新月份的工作表
            new_sheet_name = f"{roc_year}{month:02d}"
            if new_sheet_name in wb.sheetnames:
                del wb[new_sheet_name]
                
            target_sheet = wb.copy_worksheet(source_sheet)
            target_sheet.title = new_sheet_name
            
            # 更新 A1 標題中的月份
            title_text = f"臺中市馬鳴國小{roc_year}年{month}月教師 公假排代 授課節數鐘點費印領清冊"
            target_sheet['A1'] = title_text
            
            # 清除第 4 行到第 17 行的範例資料 (保留格式與公式，只清除輸入值)
            for r in range(4, 18):
                target_sheet[f'A{r}'] = None  # 姓名
                target_sheet[f'B{r}'] = None  # 節數
                target_sheet[f'C{r}'] = None  # 單價
                target_sheet[f'E{r}'] = None  # 學校勞保
                target_sheet[f'F{r}'] = None  # 學校退撫
                target_sheet[f'G{r}'] = None  # 學校健保
                target_sheet[f'J{r}'] = None  # 代扣勞保
                target_sheet[f'K{r}'] = None  # 代扣健保
                target_sheet[f'L{r}'] = None  # 代扣補充保費
                target_sheet[f'N{r}'] = None  # 備註
                # P~U 為範本右側的節數統計欄；未清除會把上一期的數字留在新清冊上
                for col in ('P', 'Q', 'R', 'S', 'T', 'U'):
                    target_sheet[f'{col}{r}'] = None

            # 按教師歸納鐘點費與假別備註
            leave_map = {
                "身心調適假": "身",
                "事假": "事",
                "病假": "病",
                "公假": "公",
                "公差假": "公差",
                "婚假": "婚",
                "喪假": "喪",
                "產假": "產",
                "其它": "其它"
            }
            
            teachers_data = {}
            for rec in period_records:
                teacher = rec.get('subTeacher')
                if not teacher:
                    continue
                if teacher not in teachers_data:
                    teachers_data[teacher] = {
                        'periods': 0,
                        'rate': rec.get('rate', 405),
                        'leaves': {}
                    }
                # 統計本筆代課中實際打勾付費的節數 (如果該筆紀錄的鐘點費為 0，代表整天代導不領鐘點費，不計入鐘點節數)
                if rec.get('classFee', 0) > 0:
                    paid_periods = sum(1 for p in rec.get('periodsDetail', []) if p.get('paid', False))
                    if paid_periods == 0:
                        paid_periods = rec.get('periodsCount', 0)
                else:
                    paid_periods = 0
                
                teachers_data[teacher]['periods'] += paid_periods
                
                # 歸類假別
                if paid_periods > 0:
                    leave_type = rec.get('leaveType', '其它')
                    short_leave = leave_map.get(leave_type, '其它')
                    teachers_data[teacher]['leaves'][short_leave] = teachers_data[teacher]['leaves'].get(short_leave, 0) + paid_periods
                
            # 將歸納好的教師資料填入 Excel
            row_idx = 4
            for teacher, t_info in teachers_data.items():
                if row_idx > 17:
                    print(f"[警告] 鐘點代課教師數量超過範本上限 (14人)，部分人員可能無法顯示！")
                    break
                    
                # 組合備註 (例如: 身6病2)
                note_parts = []
                for lv, cnt in t_info['leaves'].items():
                    if cnt > 0:
                        note_parts.append(f"{lv}{cnt}")
                note_str = "".join(note_parts)
                
                target_sheet[f'A{row_idx}'] = teacher
                target_sheet[f'B{row_idx}'] = t_info['periods']
                target_sheet[f'C{row_idx}'] = t_info['rate']
                target_sheet[f'N{row_idx}'] = note_str
                row_idx += 1
                
            print(f"成功產生鐘點代課工作表: '{new_sheet_name}' (共彙整 {len(teachers_data)} 位教師)")

        # 2. 處理整天日薪代課 (導師費代課)
        daily_records = [r for r in filtered_records if r.get('payMode', 'perPeriod') == 'mentorDaily']
        if daily_records:
            source_sheet_name = '11405 日代 印領'
            if source_sheet_name not in wb.sheetnames:
                print("[警告] 找不到日薪代課範本工作表 '11405 日代 印領'，略過日薪工作表產生。")
            else:
                print(f"使用工作表 '{source_sheet_name}' 作為日薪代課範本...")
                source_sheet = wb[source_sheet_name]
                
                new_sheet_name = f"{roc_year}{month:02d}日代"
                if new_sheet_name in wb.sheetnames:
                    del wb[new_sheet_name]
                    
                target_sheet = wb.copy_worksheet(source_sheet)
                target_sheet.title = new_sheet_name
                
                # 更新 A1 標題
                target_sheet['A1'] = f"臺中市馬鳴國小{roc_year}年{month}月教師公假排代 日薪代課 印領清冊"
                
                # 清除第 4 行到第 6 行的資料 (日薪範本只提供 3 行資料)
                for r in range(4, 7):
                    target_sheet[f'A{r}'] = None
                    target_sheet[f'B{r}'] = None
                    target_sheet[f'C{r}'] = None
                    target_sheet[f'E{r}'] = None
                    target_sheet[f'F{r}'] = None
                    target_sheet[f'G{r}'] = None
                    target_sheet[f'J{r}'] = None
                    target_sheet[f'K{r}'] = None
                    target_sheet[f'L{r}'] = None
                    target_sheet[f'N{r}'] = None
                    
                # 寫入日薪代課
                row_idx = 4
                for rec in daily_records:
                    if row_idx > 6:
                        print("[警告] 日薪代課教師數量超過範本上限 (3人)！")
                        break
                    teacher = rec.get('subTeacher')
                    days = rec.get('actingDays', 1)
                    if not days or days == 0:
                        days = 1
                        
                    print(f"\n[日薪代課] 偵測到 {teacher} 有整天代理導師紀錄。")
                    print(f"請參閱《115日薪代課教師薪水表(1).pdf》獲取適用薪水率。")
                    
                    # 預估適用日薪
                    days_in_month = 31 if month in [1, 3, 5, 7, 8, 10, 12] else 30
                    default_rate = 1399
                    if rec.get('isActingMentor', False) or rec.get('payMode') == 'mentorDaily':
                        default_rate = 1528 if days_in_month == 31 else 1532
                        
                    try:
                        rate_input = input(f"請輸入 {teacher} 的每日薪資 (預設 {default_rate} 元): ").strip()
                        daily_rate = int(rate_input) if rate_input else default_rate
                    except:
                        daily_rate = default_rate
                        
                    # 建立備註 (如: 1399(日薪)+129(導師費))
                    mentor_diff = daily_rate - 1399
                    if mentor_diff > 0:
                        note = f"1399(日薪)+{mentor_diff}(導師費)"
                    else:
                        note = "日薪代課"
                        
                    target_sheet[f'A{row_idx}'] = teacher
                    target_sheet[f'B{row_idx}'] = days
                    target_sheet[f'C{row_idx}'] = daily_rate
                    target_sheet[f'N{row_idx}'] = note
                    row_idx += 1
                    
                print(f"成功產生日薪代課工作表: '{new_sheet_name}'")
                
        # 3. 儲存檔案
        output_file = f"{roc_year}年{month:02d}月_代課費印領清冊.xlsx"
        wb.save(output_file)
        print(f"[成功] 成功產生 Excel 試算表憑證！")
        print(f"檔案名稱: {output_file}")
        print(f"存放路徑: {os.path.abspath(output_file)}")
        return True
    except Exception as e:
        print(f"[錯誤] 產生 Excel 憑證失敗: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("==================================================")
    print("       代課費印領清冊 Word 憑證生成助理           ")
    print("==================================================")
    
    # 1. 尋找目錄下的 JSON 檔案（去除重複，優先顯示本系統備份檔）
    json_files = sorted(set(glob.glob("*.json")))
    # 讓「代課費系統備份_*.json」排在最前面，方便直接按 Enter 選用
    json_files.sort(key=lambda f: (0 if "代課費" in f or "備份" in f else 1, f))
    
    if not json_files:
        print("[錯誤] 未在當前目錄下找到任何代課紀錄 JSON 檔案！")
        print("請先在網頁端點擊「備份匯出資料 (JSON)」，並將檔案放入此資料夾中。")
        input("\n按任意鍵結束...")
        return

    # 顯示檔案供選擇
    print("找到以下資料檔案：")
    for idx, f in enumerate(json_files):
        print(f" [{idx + 1}] {f}")
    
    file_choice = 0
    if len(json_files) > 1:
        try:
            choice = input(f"請選擇檔案序號 (1-{len(json_files)}, 預設 1): ").strip()
            file_choice = int(choice) - 1 if choice else 0
        except:
            file_choice = 0
    
    selected_file = json_files[file_choice]
    print(f"[選擇] 已選擇檔案: {selected_file}")

    # 讀取資料
    try:
        with open(selected_file, 'r', encoding='utf-8') as f:
            records = json.load(f)
    except Exception as e:
        print(f"[錯誤] 讀取檔案失敗: {e}")
        input("\n按任意鍵結束...")
        return

    if not isinstance(records, list) or len(records) == 0:
        print("[錯誤] 檔案內無有效的代課紀錄數據！")
        input("\n按任意鍵結束...")
        return

    # 2. 獲取所有可用的月份
    months = sorted(list(set(r['date'][:7] for r in records)))
    print("\n可結算月份：")
    for idx, m in enumerate(months):
        print(f" [{idx + 1}] {m}")
    
    month_choice = 0
    if len(months) > 1:
        try:
            choice = input(f"請選擇結算月份序號 (1-{len(months)}, 預設最後一個月): ").strip()
            month_choice = int(choice) - 1 if choice else len(months) - 1
        except:
            month_choice = len(months) - 1
    else:
        month_choice = 0
        
    target_month = months[month_choice]
    print(f"[選擇] 已選擇結算月份: {target_month}")

    # 篩選該月份資料
    filtered_records = [r for r in records if r['date'].startswith(target_month)]
    # 按日期排序
    filtered_records.sort(key=lambda x: x['date'])

    if not filtered_records:
        print(f"[錯誤] 該月份 ({target_month}) 沒有任何代課紀錄！")
        input("\n按任意鍵結束...")
        return

    # 3. 開始產生 Word 文件
    doc = docx.Document()
    
    # 設定頁面為橫向 A4
    section = doc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    # 交換寬高以達到橫向
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    # 設定邊界 (各 0.8 英吋)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # 設定預設字型
    style = doc.styles['Normal']
    font = style.font
    font.name = 'DFKai-SB'  # 標楷體
    font.size = Pt(11)
    
    # 標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"台中市馬鳴國民小學 {target_month.split('-')[0]} 年 {target_month.split('-')[1]} 月份 外聘教師代課鐘點費及代理導師費印領清冊")
    run.font.size = Pt(16)
    run.font.bold = True
    
    # 建立表格
    # 欄位：日期, 班級, 請假教師, 代課教師, 代課節數, 單價, 鐘點費小計, 代理導師天數, 代理導師費, 應領總額, 領款人簽章
    headers = ["代課日期", "班級", "請假教師", "代課教師", "代課節數", "鐘點費單價", "鐘點費小計", "代理導師天數", "代理導師費", "應領總額", "領款人簽章\n(或蓋章)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    # 設定標頭
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "E6E6E6")  # 灰色背景
        # 置中對齊
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10.5)

    # 填入資料
    total_periods = 0
    total_class_fee = 0
    total_mentor_days = 0
    total_mentor_fee = 0
    grand_total = 0

    border_format = {'sz': 4, 'val': 'single', 'color': '000000'}

    has_pending = False  # 是否有金額待人工填寫的整天代理紀錄

    for rec in filtered_records:
        row_cells = table.add_row().cells

        # 整天代理（mentorDaily）：教師端不送天數與金額，設計上由行政人員依日薪表計算。
        # 這類紀錄不領鐘點費，但「代理導師天數」應為 1 天，金額須人工填寫，
        # 不可直接印 $0 —— 否則整天代理會在清冊上消失。
        is_daily = rec.get('payMode') == 'mentorDaily' or (
            rec.get('isActingMentor', False) and rec.get('classFee', 0) == 0
        )

        periods_printed = rec.get('periodsCount', 0) if rec.get('classFee', 0) > 0 else 0
        acting_days = rec.get('actingDays', 0)
        if is_daily and acting_days == 0:
            acting_days = 1  # 一筆整天代理 = 當天 1 天

        mentor_fee = rec.get('mentorFee', 0)
        total_fee = rec.get('totalFee', 0)
        pending = is_daily and mentor_fee == 0 and total_fee == 0
        if pending:
            has_pending = True

        total_periods += periods_printed
        total_class_fee += rec.get('classFee', 0)
        total_mentor_days += acting_days
        total_mentor_fee += mentor_fee
        grand_total += total_fee

        # 填值 (強制轉為 str 以防數字型別觸發 python-docx 的 iterable 錯誤)
        row_cells[0].text = str(format_roc_date(rec.get('date', '')) or '')
        row_cells[1].text = str(rec.get('className', '') or '')
        row_cells[2].text = str(rec.get('absentTeacher', '') or '')
        row_cells[3].text = str(rec.get('subTeacher', '') or '')
        row_cells[4].text = str(periods_printed)
        row_cells[5].text = "" if is_daily else str(rec.get('rate', 405))
        row_cells[6].text = f"${int(rec.get('classFee', 0) or 0):,}"
        row_cells[7].text = str(acting_days)
        row_cells[8].text = "待填(日薪)" if pending else f"${int(mentor_fee or 0):,}"
        row_cells[9].text = "待填(日薪)" if pending else f"${int(total_fee or 0):,}"
        row_cells[10].text = "" # 簽章欄留空

        # 對齊與邊框
        for i, cell in enumerate(row_cells):
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            p = cell.paragraphs[0]
            if i in [0, 1, 2, 3, 10]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in [4, 5, 7]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 新增合計列
    row_cells = table.add_row().cells
    row_cells[0].text = "合  計"
    # 合併前4格
    row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3])
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    
    row_cells[4].text = str(total_periods)
    row_cells[5].text = ""
    row_cells[6].text = f"${total_class_fee:,}"
    row_cells[7].text = f"{total_mentor_days:g}" if total_mentor_days > 0 else "0"
    # 有整天代理待填時，合計不可印成看似完整的金額
    row_cells[8].text = f"${total_mentor_fee:,}＋待填" if has_pending else f"${total_mentor_fee:,}"
    row_cells[9].text = f"${grand_total:,}＋待填" if has_pending else f"${grand_total:,}"
    row_cells[10].text = ""

    # 設定合計列的對齊與邊框
    for i, cell in enumerate(row_cells):
        set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
        set_cell_background(cell, "F2F2F2")
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].font.bold = True
        if i == 0:
            pass
        elif i in [4, 5, 7]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i in [6, 8, 9]:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 設定表格外觀寬度與微調
    for row in table.rows:
        for i, w in enumerate([1.2, 0.6, 0.8, 0.8, 0.6, 0.8, 1.0, 1.0, 1.0, 1.1, 1.3]):
            row.cells[i].width = Inches(w)

    # 有整天代理待填金額時，於表格下方明確提示，避免被誤認為已結算完成
    if has_pending:
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(
            "※ 標示「待填(日薪)」者為整天代理導師，係按日薪計算，"
            "請依《日薪代課教師薪水表》核算後補填金額，本表金額尚未完整。"
        )
        note_run.font.size = Pt(10)
        note_run.font.bold = True

    # 新增簽章欄位
    doc.add_paragraph() # 空行
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(20)
    
    # 增加簽章文字
    t_run1 = footer_p.add_run("承辦人：　　　　　　　　　　")
    t_run1.font.size = Pt(12)
    t_run1.font.bold = True
    
    t_run2 = footer_p.add_run("主計主任：　　　　　　　　　　")
    t_run2.font.size = Pt(12)
    t_run2.font.bold = True
    
    t_run3 = footer_p.add_run("校長：　　　　　　　　　　")
    t_run3.font.size = Pt(12)
    t_run3.font.bold = True

    # 存檔
    output_filename = f"{target_month.replace('-', '年')}月_代課費印領清冊.docx"
    try:
        doc.save(output_filename)
        print(f"\n[成功] 成功產生 Word 憑證！")
        print(f"檔案名稱: {output_filename}")
        print(f"存放路徑: {os.path.abspath(output_filename)}")
    except Exception as e:
        print(f"[錯誤] 儲存 Word 檔案失敗: {e}")

    # 產生 Excel 憑證 (套用學校 Excel 印領清冊格式)
    generate_excel(target_month, filtered_records)
        
    input("\n處理完成，按任意鍵結束...")

if __name__ == "__main__":
    main()
